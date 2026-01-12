import os
import sys
import glob
import uuid
import pandas as pd
import yaml
import streamlit as st
from langchain_community.document_loaders import (
    DirectoryLoader, 
    PyPDFLoader
)
from langchain_text_splitters import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from sentence_transformers import CrossEncoder
from langchain_core.documents import Document
from rank_bm25 import BM25Okapi
from langchain_ollama import ChatOllama
from docx import Document as DocxDocument

# 设置环境变量
os.environ["NO_PROXY"] = "localhost,127.0.0.1"

# 添加 src 目录到路径
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "src"))

# 导入改进模块
from query_decomposition import QueryDecomposer
from hyde_retrieval import HyDERetriever
from rrf_fusion import rrf_fuse
from pdf_table_extractor import PDFTableExtractor, load_pdfs_with_table_extraction

# --- Config Loader ---
def load_config():
    try:
        with open("config.yaml", "r", encoding="utf-8") as f:
            return yaml.safe_load(f)
    except:
        st.error("Missing config.yaml!")
        return {}

CONFIG = load_config()

# --- 数据加载函数 ---

def load_excel_as_text(directory):
    documents = []
    excel_files = glob.glob(os.path.join(directory, "**/*.xlsx"), recursive=True) + \
                  glob.glob(os.path.join(directory, "**/*.csv"), recursive=True)
                  
    for file_path in excel_files:
        try:
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            df = df.fillna("Unknown")
            columns = df.columns.tolist()
            for index, row in df.iterrows():
                content_parts = [f"Source File: {os.path.basename(file_path)}"]
                for col in columns:
                    val = row.get(col, '')
                    content_parts.append(f"{col}: {val}")
                content = " | ".join(content_parts)
                doc = Document(
                    page_content=content, 
                    metadata={"source": file_path, "type": "excel_record"}
                )
                documents.append(doc)
        except Exception as e:
            st.error(f"表格加载失败 {file_path}: {e}")
    return documents

def load_word_with_structure(directory):
    documents = []
    word_files = glob.glob(os.path.join(directory, "**/*.docx"), recursive=True)
    for file_path in word_files:
        try:
            doc_obj = DocxDocument(file_path)
            current_heading = "General"
            for para in doc_obj.paragraphs:
                text = para.text.strip()
                if not text: continue
                if 'Heading' in para.style.name:
                    current_heading = text
                else:
                    enhanced_content = f"Document Section [{current_heading}]: {text}"
                    doc = Document(
                        page_content=enhanced_content, 
                        metadata={"source": file_path, "type": "contract_clause", "section": current_heading}
                    )
                    documents.append(doc)
        except Exception as e:
            st.error(f"Word 加载失败 {file_path}: {e}")
    return documents

def load_images_with_vlm(directory):
    documents = []
    image_files = glob.glob(os.path.join(directory, "**/*.png"), recursive=True) + \
                  glob.glob(os.path.join(directory, "**/*.jpg"), recursive=True)
    
    # 从 Config 读取模型名称
    vlm_name = CONFIG["system"].get("vlm_model_name", "llava")
    
    try:
        vlm = ChatOllama(model=vlm_name, base_url="http://127.0.0.1:11434")
        vlm.invoke("hi")
        has_vlm = True
    except:
        has_vlm = False
        st.warning(f"⚠️ 未检测到本地 {vlm_name} 模型，将使用模拟描述处理图片。")

    for file_path in image_files:
        try:
            if has_vlm:
                from langchain_core.messages import HumanMessage
                import base64
                def encode_image(image_path):
                    with open(image_path, "rb") as image_file:
                        return base64.b64encode(image_file.read()).decode('utf-8')
                base64_image = encode_image(file_path)
                prompt = "Please describe this image in detail. Extract any text, numbers, or technical specifications visible in the image."
                msg = HumanMessage(content=[
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                ])
                response = vlm.invoke([msg])
                description = response.content
                st.info(f"🖼️ [VLM] 已分析图片: {os.path.basename(file_path)}")
            else:
                # 依然保留 Fallback 以防万一
                description = f"Image file: {os.path.basename(file_path)}"
                
            doc = Document(
                page_content=f"Image Description [{os.path.basename(file_path)}]: {description}", 
                metadata={"source": file_path, "type": "image_caption", "image_path": file_path}
            )
            documents.append(doc)
        except Exception as e:
            st.error(f"图片加载失败 {file_path}: {e}")
    return documents

@st.cache_resource
def initialize_rag_system():
    status = st.empty()
    status.info("正在初始化 RAG 系统，这可能需要几十秒...")

    # 1. 加载数据
    docs = []
    # Markdown
    md_files = glob.glob(os.path.join("knowledge_base", "**/*.md"), recursive=True) + \
               glob.glob(os.path.join("downloads", "**/*.md"), recursive=True)
    for f in md_files:
        try:
            with open(f, 'r', encoding='utf-8') as file:
                md_content = file.read()
            headers_to_split_on = [("#", "Header 1"), ("##", "Header 2"), ("###", "Header 3")]
            markdown_splitter = MarkdownHeaderTextSplitter(headers_to_split_on=headers_to_split_on)
            md_header_splits = markdown_splitter.split_text(md_content)
            for doc in md_header_splits:
                doc.metadata["source"] = f
                doc.metadata["type"] = "markdown_section"
                header_path = " > ".join([v for k, v in doc.metadata.items() if k.startswith("Header")])
                
                content = doc.page_content
                keywords_failure = ["failure", "incident", "lost pulse", "stuck", "fishing", "pooh"]
                if "report" in f.lower() or "log" in f.lower():
                    if any(kw in content.lower() for kw in keywords_failure):
                        content = f"[EVENT: Drilling Stopped / Operational Interruption / NPT]\n{content}"
                
                if header_path:
                    doc.page_content = f"Section [{header_path}]:\n{content}"
                else:
                    doc.page_content = content 
            docs.extend(md_header_splits)
        except Exception as e:
            st.error(f"Markdown 加载失败 {f}: {e}")

    docs.extend(load_word_with_structure("downloads"))
    docs.extend(load_excel_as_text("downloads"))
    docs.extend(load_images_with_vlm("downloads"))
    
    # [优化] 使用增强版 PDF 加载器，自动提取表格
    st.info("📄 正在提取 PDF 表格...")
    pdf_docs = load_pdfs_with_table_extraction("downloads")
    docs.extend(pdf_docs)
    st.success(f"✅ PDF 加载完成: {len(pdf_docs)} 个文档片段")

    # 2. Parent-Child Splitting (使用 Config 参数)
    p_size = CONFIG["indexing"]["chunk_size_parent"]
    p_overlap = CONFIG["indexing"]["chunk_overlap_parent"]
    c_size = CONFIG["indexing"]["chunk_size_child"]
    c_overlap = CONFIG["indexing"]["chunk_overlap_child"]
    
    parent_splitter = RecursiveCharacterTextSplitter(chunk_size=p_size, chunk_overlap=p_overlap)
    child_splitter = RecursiveCharacterTextSplitter(chunk_size=c_size, chunk_overlap=c_overlap)
    
    parent_docs = []
    child_docs = []
    docstore = {} 
    
    for raw_doc in docs:
        if len(raw_doc.page_content) > p_size:
            parents = parent_splitter.split_documents([raw_doc])
        else:
            parents = [raw_doc]
            
        for parent in parents:
            parent_id = str(uuid.uuid4())
            parent.metadata["doc_id"] = parent_id
            docstore[parent_id] = parent
            parent_docs.append(parent)
            
            children = child_splitter.split_documents([parent])
            for child in children:
                child.metadata["parent_id"] = parent_id
                parent_content = parent.page_content
                if "[EVENT:" in parent_content:
                    tag_line = parent_content.split("\n")[0]
                    if "[EVENT:" in tag_line and tag_line not in child.page_content:
                        child.page_content = f"{tag_line}\n{child.page_content}"
                child_docs.append(child)
                
    splits = child_docs
    print(f"DEBUG: Created {len(parent_docs)} Parents and {len(child_docs)} Children.")

    # 3. VectorStore
    emb_model = CONFIG["system"]["embedding_model"]
    embeddings = HuggingFaceEmbeddings(model_name=emb_model)
    vectorstore = Chroma.from_documents(documents=splits, embedding=embeddings)

    # 4. BM25
    tokenized_corpus = [doc.page_content.split(" ") for doc in splits]
    bm25 = BM25Okapi(tokenized_corpus)

    # 5. Reranker
    rerank_model = CONFIG["system"]["reranker_model"]
    reranker = CrossEncoder(rerank_model)

    # 6. LLM
    llm_name = CONFIG["system"]["model_name"]
    try:
        llm = ChatOllama(model=llm_name, base_url="http://127.0.0.1:11434")
        llm.invoke("Hi")
    except Exception as e:
        st.error(f"本地 LLM 初始化失败: {e}")
        llm = None

    status.success("RAG 系统初始化完成 (Config Loaded)！")
    return {
        "vectorstore": vectorstore,
        "bm25": bm25,
        "reranker": reranker,
        "llm": llm,
        "splits": splits,
        "docstore": docstore
    }

# --- CRAG Components (使用 Config Prompt) ---

def grade_documents(query, documents, llm):
    if not llm or not documents:
        return documents, False
    
    relevant_docs = []
    
    # 从 Config 读取 Prompt
    prompt_template = CONFIG["prompts"]["grade"]
    
    has_relevant = False
    with st.status("👩‍🏫 正在评估文档质量...", expanded=False) as status:
        for doc in documents:
            prompt = f"{prompt_template}\n\nQuestion: {query}\nDocument: {doc.page_content[:800]}"
            try:
                score = llm.invoke(prompt).content.strip().lower()
                doc_name = os.path.basename(doc.metadata.get("source", "unknown"))
                if "yes" in score:
                    relevant_docs.append(doc)
                    has_relevant = True
                    status.write(f"✅ 相关: {doc_name}")
                else:
                    status.write(f"🗑️ 过滤: {doc_name}")
            except:
                relevant_docs.append(doc)
    
    needs_rewrite = (len(relevant_docs) == 0) or (len(relevant_docs) < len(documents) / 2)
    return relevant_docs, needs_rewrite

def rewrite_query(original_query, llm):
    # 从 Config 读取 Prompt
    prompt_template = CONFIG["prompts"]["rewrite"]
    prompt = prompt_template.format(original_query=original_query)
    
    try:
        response = llm.invoke(prompt)
        return response.content.strip()
    except:
        return original_query

def _retrieve_documents(query, vectorstore, bm25, reranker, splits, docstore, llm=None):
    """
    混合检索函数 - 使用 RRF 融合多路检索结果
    
    检索策略:
    1. 多路检索（分组向量 + BM25 + HyDE）
    2. [NEW] RRF 融合（比简单加权更鲁棒）
    3. Parent Document 还原
    4. Reranker 精排
    """
    debug_info = {"retrieval_sources": []}
    retrieval_conf = CONFIG["retrieval"]
    
    # ========== 收集多路检索结果（保持排序） ==========
    all_rankings = []  # [(docs_list, weight, name), ...]
    
    # ----- 路径1: 分组向量检索 -----
    filters = [
        {"name": "Excel/CSV", "filter": {"type": "excel_record"}, "k": retrieval_conf["k_excel"], "weight": 1.0}, 
        {"name": "Word", "filter": {"type": "contract_clause"}, "k": retrieval_conf["k_word"], "weight": 1.0},
        {"name": "Image", "filter": {"type": "image_caption"}, "k": retrieval_conf["k_image"], "weight": 1.0},
        {"name": "Technical", "filter": {"type": "markdown_section"}, "k": retrieval_conf["k_tech"], "weight": 1.0}, 
        {"name": "General", "filter": None, "k": retrieval_conf["k_general"], "weight": 0.8}  # General 权重略低
    ]
    
    for f in filters:
        effective_k = f["k"]
        # 动态调整：财务问题时提升 Excel 权重
        if f["name"] == "Excel/CSV" and any(kw in query.lower() for kw in ["growth", "revenue", "%", "billion"]):
            effective_k = retrieval_conf["k_financial_boost"]
            f["weight"] = 1.5  # 提升权重
            
        kwargs = {"k": effective_k}
        if f["filter"]: 
            kwargs["filter"] = f["filter"]
        try:
            sub_docs = vectorstore.similarity_search(query, **kwargs)
            # 标记 boost
            for doc in sub_docs:
                if f["name"] == "Image" and "image" in query.lower():
                    doc.metadata["boost"] = True
                doc.metadata["retrieval_method"] = f"grouped_{f['name'].lower()}"
            
            if sub_docs:
                all_rankings.append((sub_docs, f["weight"], f"vector_{f['name']}"))
                debug_info["retrieval_sources"].append(f"vector_{f['name']}: {len(sub_docs)}")
        except: 
            pass
    
    # ----- 路径2: BM25 关键词检索 -----
    tokenized_query = query.split(" ")
    bm25_docs = bm25.get_top_n(tokenized_query, splits, n=15)
    for doc in bm25_docs:
        doc.metadata["retrieval_method"] = "bm25"
    if bm25_docs:
        all_rankings.append((bm25_docs, 0.8, "bm25"))  # BM25 权重
        debug_info["retrieval_sources"].append(f"bm25: {len(bm25_docs)}")
    
    # ----- 路径3: HyDE 检索 -----
    hyde_debug = None
    if llm:
        try:
            hyde_retriever = HyDERetriever(llm, vectorstore, reranker=None)
            hyde_docs, hyde_debug = hyde_retriever.retrieve(
                query, 
                k=10,
                use_hyde=True,
                combine_with_original=False
            )
            for doc in hyde_docs:
                doc.metadata["retrieval_method"] = "hyde"
            if hyde_docs:
                all_rankings.append((hyde_docs, 1.2, "hyde"))  # HyDE 权重较高
                debug_info["retrieval_sources"].append(f"hyde: {len(hyde_docs)}")
            if hyde_debug:
                debug_info["hyde"] = hyde_debug
        except Exception as e:
            debug_info["hyde_error"] = str(e)
    
    if not all_rankings:
        return [], debug_info

    # ========== RRF 融合 ==========
    rankings = [r[0] for r in all_rankings]
    weights = [r[1] for r in all_rankings]
    
    fused_results = rrf_fuse(
        rankings=rankings,
        weights=weights,
        k=60,  # RRF 常数
        top_n=20  # 多取一些，后面还要 rerank
    )
    
    candidate_child_docs = [doc for doc, score in fused_results]
    debug_info["rrf_fusion"] = f"Fused {len(rankings)} rankings -> {len(candidate_child_docs)} docs"

    # ========== Parent Document 还原 ==========
    candidate_parents = []
    seen_parent_ids = set()
    for child in candidate_child_docs:
        parent_id = child.metadata.get("parent_id")
        if parent_id and parent_id in docstore:
            if parent_id not in seen_parent_ids:
                parent_doc = docstore[parent_id]
                if child.metadata.get("boost"): 
                    parent_doc.metadata["boost"] = True
                parent_doc.metadata["retrieval_method"] = child.metadata.get("retrieval_method", "unknown")
                candidate_parents.append(parent_doc)
                seen_parent_ids.add(parent_id)
        else:
            candidate_parents.append(child)

    if not candidate_parents: 
        return [], debug_info

    # ========== Reranker 精排 ==========
    pairs = [[query, doc.page_content] for doc in candidate_parents]
    scores = reranker.predict(pairs)
    doc_score_pairs = list(zip(candidate_parents, scores))
    
    final_pairs = []
    for doc, score in doc_score_pairs:
        final_score = score
        # 额外加分规则
        if doc.metadata.get("boost"): 
            final_score += 0.5
        if "[EVENT:" in doc.page_content and any(q in query.lower() for q in ["why", "reason", "stop", "fail"]):
            final_score += 2.0
        final_pairs.append((doc, final_score))
        
    final_pairs.sort(key=lambda x: x[1], reverse=True)
    
    return [doc for doc, score in final_pairs[:5]], debug_info

def process_query(query, rag_components):
    """
    处理用户查询 - 支持 Query Decomposition
    
    流程:
    1. [NEW] 问题分解：判断是否需要拆分为子问题
    2. 检索：对每个子问题分别检索
    3. CRAG：评估文档质量，必要时重写查询
    4. 聚合：合并所有检索结果
    5. 生成：基于聚合的上下文生成回答
    """
    vectorstore = rag_components["vectorstore"]
    bm25 = rag_components["bm25"]
    reranker = rag_components["reranker"]
    llm = rag_components["llm"]
    splits = rag_components["splits"]
    docstore = rag_components["docstore"]
    
    debug_info = []
    
    # ============ Step 1: Query Decomposition ============
    decomposer = QueryDecomposer(llm)
    decomposition = decomposer.decompose(query)
    
    is_complex = decomposition["is_complex"]
    sub_queries = decomposition["sub_queries"]
    aggregation_type = decomposition["aggregation_type"]
    
    if is_complex:
        st.info(f"🔍 检测到复杂问题，分解为 {len(sub_queries)} 个子问题")
        debug_info.append(f"Query Decomposed: {sub_queries}")
        debug_info.append(f"Aggregation Type: {aggregation_type}")
    
    # ============ Step 2: 分别检索每个子问题 ============
    all_docs = []
    sub_results = []  # 记录每个子问题的检索结果
    seen_contents = set()
    
    for i, sub_q in enumerate(sub_queries):
        if is_complex:
            with st.status(f"🔍 检索子问题 {i+1}/{len(sub_queries)}: {sub_q[:50]}...", expanded=False):
                docs, retrieve_debug = _retrieve_documents(sub_q, vectorstore, bm25, reranker, splits, docstore, llm=llm)
        else:
            docs, retrieve_debug = _retrieve_documents(sub_q, vectorstore, bm25, reranker, splits, docstore, llm=llm)
        
        # 记录 HyDE 调试信息
        if retrieve_debug.get("hyde"):
            debug_info.append(f"HyDE: {retrieve_debug['hyde'].get('hypothetical_doc', 'N/A')[:100]}...")
        
        # 去重并记录
        unique_docs = []
        for doc in docs:
            if doc.page_content not in seen_contents:
                unique_docs.append(doc)
                seen_contents.add(doc.page_content)
                all_docs.append(doc)
        
        # 记录子问题结果（用于聚合生成）
        context_preview = "\n".join([d.page_content[:300] for d in unique_docs[:2]])
        sub_results.append({
            "sub_query": sub_q,
            "context": context_preview,
            "docs": unique_docs
        })
    
    # ============ Step 3: CRAG 自我修正 ============
    if not llm:
        final_docs = all_docs[:5]
        debug_info.append("CRAG skipped (No LLM)")
    else:
        # 对聚合后的文档进行质量评估
        filtered_docs, needs_rewrite = grade_documents(query, all_docs[:6], llm)
        
        if needs_rewrite and not is_complex:
            # 只对简单问题启用 CRAG 重写（复杂问题已经分解过了）
            st.warning("⚠️ 检索质量预警：正在尝试自我修正...")
            better_query = rewrite_query(query, llm)
            st.info(f"🔄 修正后查询: {better_query}")
            retry_docs, _ = _retrieve_documents(better_query, vectorstore, bm25, reranker, splits, docstore, llm=llm)
            final_docs = filtered_docs + retry_docs
            seen = set()
            unique_docs = []
            for d in final_docs:
                if d.page_content not in seen:
                    unique_docs.append(d)
                    seen.add(d.page_content)
            final_docs = unique_docs[:5]
            debug_info.append(f"CRAG Activated. Rewrote to: {better_query}")
        else:
            final_docs = filtered_docs if filtered_docs else all_docs[:5]
            debug_info.append("CRAG Passed" if filtered_docs else "CRAG: Using original docs")

    # ============ Step 4: 生成回答 ============
    if llm and final_docs:
        if is_complex:
            # 复杂问题：使用聚合 Prompt
            sub_results_text = ""
            for i, sr in enumerate(sub_results):
                sub_results_text += f"\n--- Sub-question {i+1}: {sr['sub_query']} ---\n"
                sub_results_text += f"Context:\n{sr['context']}\n"
            
            agg_prompt = decomposer.get_aggregation_prompt(aggregation_type)
            prompt = agg_prompt.format(
                sub_results=sub_results_text,
                original_query=query
            )
        else:
            # 简单问题：使用原有 Prompt
            context = "\n\n".join([d.page_content for d in final_docs])
            prompt_template = CONFIG["prompts"]["generation"]
            prompt = prompt_template.format(context=context, query=query)
        
        response_obj = llm.invoke(prompt)
        response = response_obj.content
        final_pairs = [(d, 0.99) for d in final_docs]
    else:
        response = "No relevant documents found or LLM offline."
        final_pairs = []

    return response, final_pairs, debug_info

# --- UI Setup ---
st.set_page_config(page_title="OilField RAG Assistant (Configurable)", layout="wide")
st.title("🛢️ OilField RAG Assistant (Competition Ready)")
st.markdown("本地油田服务领域知识库问答系统 - 5人协作版")

if "messages" not in st.session_state:
    st.session_state.messages = []

rag = initialize_rag_system()

with st.sidebar:
    st.header("系统状态")
    if rag["llm"]:
        st.success("🟢 LLM Online")
        st.info(f"Using Model: {CONFIG['system']['model_name']}")
    else:
        st.error("🔴 LLM Offline")

for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])
        if "sources" in message:
            with st.expander("查看参考文档"):
                for i, (doc, _) in enumerate(message["sources"]):
                    st.markdown(f"**Doc {i+1}: {os.path.basename(doc.metadata.get('source', 'unknown'))}**")
                    st.text(doc.page_content[:300] + "...")

if prompt := st.chat_input("请输入问题..."):
    st.chat_message("user").markdown(prompt)
    st.session_state.messages.append({"role": "user", "content": prompt})

    with st.chat_message("assistant"):
        with st.spinner("思考与检索中..."):
            response, sources, debug_info = process_query(prompt, rag)
            st.markdown(response)
            with st.expander("调试信息 & 参考文档"):
                st.write(debug_info)
                for i, (doc, _) in enumerate(sources):
                    st.markdown(f"**Doc {i+1}: {os.path.basename(doc.metadata.get('source', 'unknown'))}**")
                    if doc.metadata.get("type") == "image_caption" and "image_path" in doc.metadata:
                         st.image(doc.metadata["image_path"], width=300)
                    st.text(doc.page_content)
    
    st.session_state.messages.append({"role": "assistant", "content": response, "sources": sources})
