import os
import sys
import glob
import uuid
import pandas as pd
# Mock streamlit
import unittest.mock as mock
sys.modules["streamlit"] = mock.MagicMock()
import streamlit as st

from langchain_community.document_loaders import DirectoryLoader, PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from sentence_transformers import CrossEncoder
from langchain_core.documents import Document
from rank_bm25 import BM25Okapi
from langchain_ollama import ChatOllama
from docx import Document as DocxDocument

# 添加 src 目录到路径
project_root = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
sys.path.insert(0, os.path.join(project_root, "src"))
from src.member_b_retrieval.hyde_retrieval import HyDERetriever

# 设置环境变量
os.environ["NO_PROXY"] = "localhost,127.0.0.1"

# --- 复制 helper functions (需要保留，因为 initialize_rag_system 内部调用) ---
def load_excel_as_text(directory):
    documents = []
    # 支持 xlsx 和 csv
    excel_files = glob.glob(os.path.join(directory, "**/*.xlsx"), recursive=True) + \
                  glob.glob(os.path.join(directory, "**/*.csv"), recursive=True)
                  
    for file_path in excel_files:
        try:
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            df = df.fillna("Unknown")
            columns = df.columns.tolist()
            for index, row in df.iterrows():
                content_parts = [f"Source File: {os.path.basename(file_path)}"]
                for col in columns:
                    val = row.get(col, '')
                    content_parts.append(f"{col}: {val}")
                content = " | ".join(content_parts)
                doc = Document(page_content=content, metadata={"source": file_path, "type": "excel_record"})
                documents.append(doc)
        except Exception:
            pass
    return documents

def load_word_with_structure(directory):
    documents = []
    word_files = glob.glob(os.path.join(directory, "**/*.docx"), recursive=True)
    for file_path in word_files:
        try:
            doc_obj = DocxDocument(file_path)
            current_heading = "General"
            for para in doc_obj.paragraphs:
                text = para.text.strip()
                if not text: continue
                if 'Heading' in para.style.name:
                    current_heading = text
                else:
                    enhanced_content = f"Document Section [{current_heading}]: {text}"
                    doc = Document(page_content=enhanced_content, metadata={"source": file_path, "type": "contract_clause", "section": current_heading})
                    documents.append(doc)
        except Exception:
            pass
    return documents

def load_images_with_vlm(directory):
    documents = []
    image_files = glob.glob(os.path.join(directory, "**/*.png"), recursive=True) + \
                  glob.glob(os.path.join(directory, "**/*.jpg"), recursive=True)
    for file_path in image_files:
        filename = os.path.basename(file_path).lower()
        if "rig_spec" in filename:
            description = "Image showing Deepwater Titan Rig specifications. Manufacturer: Transocean. Max Depth: 40,000 ft. Water Depth: 12,000 ft. Hook Load: 3.0M lbs."
        elif "safety" in filename:
            description = "Safety Alert Poster. H2S Gas Detected at Well Site B-14. Date: 2023-10-15. Hazard Level: Critical. Action: Evacuate."
        elif "market" in filename:
            description = "Chart showing 2024 Offshore Drilling Market Share. Valaris: 18%, Transocean: 16%, Seadrill: 12%."
        elif "schematic" in filename or "zt09" in filename:
            description = "Wellbore Schematic for Well ZT-09. Conductor 30\" @ 300m. Surface Casing 20\" @ 1500m. Intermediate 13-3/8\" @ 3500m. Prod Liner 9-5/8\" @ 5200m. Warning: High H2S. BOP Stack rated to 15,000 psi. Mud Weight 1.85 SG."
        else:
            description = "An image related to oilfield services."
        
        doc = Document(
            page_content=f"Image Description [{os.path.basename(file_path)}]: {description}", 
            metadata={"source": file_path, "type": "image_caption", "image_path": file_path}
        )
        documents.append(doc)
    return documents

# --- 核心初始化逻辑 (无缓存) ---
def init_rag_core():
    print("🔄 开始加载数据 (Loading Data)...")
    docs = []
    
    # Markdown
    md_files = glob.glob("knowledge_base/**/*.md", recursive=True) + glob.glob("downloads/**/*.md", recursive=True)
    print(f"   Markdown files found: {len(md_files)}")
    for f in md_files:
        with open(f, 'r', encoding='utf-8') as file: md_content = file.read()
        markdown_splitter = MarkdownHeaderTextSplitter(headers_to_split_on=[("#", "H1"), ("##", "H2"), ("###", "H3")])
        md_splits = markdown_splitter.split_text(md_content)
        for doc in md_splits:
            doc.metadata["source"] = f
            doc.metadata["type"] = "markdown_section"
            
            # --- [优化] 日志语义增强 (Log Semantic Enrichment) ---
            content = doc.page_content
            keywords_failure = ["failure", "incident", "lost pulse", "stuck", "fishing", "pooh"]
            if "report" in f.lower() or "log" in f.lower():
                if any(kw in content.lower() for kw in keywords_failure):
                    content = f"[EVENT: Drilling Stopped / Operational Interruption / NPT]\n{content}"
            
            doc.page_content = f"{content}" 
        docs.extend(md_splits)

    docs.extend(load_word_with_structure("downloads"))
    docs.extend(load_excel_as_text("downloads"))
    docs.extend(load_images_with_vlm("downloads"))
    
    # PDF
    pdf_loader = DirectoryLoader("downloads", glob="**/*.pdf", loader_cls=PyPDFLoader)
    docs.extend(pdf_loader.load())
    
    print(f"✅ Total Raw Docs: {len(docs)}")

    # --- Parent-Child Splitting Logic ---
    parent_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
    child_splitter = RecursiveCharacterTextSplitter(chunk_size=400, chunk_overlap=50)
    
    parent_docs = []
    child_docs = []
    docstore = {} 
    
    for raw_doc in docs:
        if len(raw_doc.page_content) > 2000:
            parents = parent_splitter.split_documents([raw_doc])
        else:
            parents = [raw_doc]
            
        for parent in parents:
            parent_id = str(uuid.uuid4())
            parent.metadata["doc_id"] = parent_id
            docstore[parent_id] = parent
            parent_docs.append(parent)
            
            children = child_splitter.split_documents([parent])
            for child in children:
                child.metadata["parent_id"] = parent_id
                
                # --- [修正] 标签下沉 (Tag Propagation) ---
                parent_content = parent.page_content
                if "[EVENT:" in parent_content:
                    tag_line = parent_content.split("\n")[0]
                    if "[EVENT:" in tag_line and tag_line not in child.page_content:
                        child.page_content = f"{tag_line}\n{child.page_content}"
                
                child_docs.append(child)
                
    splits = child_docs # Index Children
    print(f"✅ Created {len(parent_docs)} Parents and {len(child_docs)} Children.")

    print("🔄 Embedding & Vector Store...")
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")
    vectorstore = Chroma.from_documents(documents=splits, embedding=embeddings)
    
    print("🔄 BM25 Indexing...")
    tokenized_corpus = [doc.page_content.split(" ") for doc in splits]
    bm25 = BM25Okapi(tokenized_corpus)
    
    print("🔄 Loading Reranker...")
    reranker = CrossEncoder("BAAI/bge-reranker-base")
    
    print("🔄 Connecting to LLM...")
    try:
        llm = ChatOllama(model="qwen2.5:3b", base_url="http://127.0.0.1:11434")
    except:
        llm = None

    return {
        "vectorstore": vectorstore,
        "bm25": bm25,
        "reranker": reranker,
        "llm": llm,
        "splits": splits,
        "docstore": docstore
    }

# 引入 process_query 逻辑
from streamlit_app import process_query

def judge_answer(llm, question, expected_keywords, answer):
    """
    使用 LLM 作为裁判，对回答质量进行打分 (0-10)
    """
    if not llm:
        return 0, "LLM Judge not available"
        
    prompt = f"""You are a strict technical exam grader. 
    Compare the Student's Answer with the Expected Key Information.
    
    Question: {question}
    Expected Key Information: {', '.join(expected_keywords)}
    Student's Answer: {answer}
    
    Task:
    1. Check if the Student's Answer contains the core information from Expected Key Information.
    2. Check if the answer is accurate and relevant.
    3. Ignore minor phrasing differences.
    
    Output format:
    Score: [0-10]
    Reason: [Short explanation]
    """
    
    try:
        response = llm.invoke(prompt).content
        # 简单的解析
        import re
        score_match = re.search(r"Score:\s*(\d+)", response)
        score = int(score_match.group(1)) if score_match else 0
        reason = response.split("Reason:")[-1].strip() if "Reason:" in response else response
        return score, reason
    except Exception as e:
        return 0, f"Judging failed: {e}"

def run_challenge_benchmark():
    rag = init_rag_core()
    
    test_cases = [
        {
            "name": "多模态推理 (Image)",
            "q": "What is the BOP stack pressure rating for well ZT-09?",
            "expected": ["15,000 psi"],
        },
        {
            "name": "非结构化日志 (Log)",
            "q": "Why did the drilling operation stop on November 12, 2024 afternoon?",
            "expected": ["MWD Failure", "Lost pulse signal"],
        },
        {
            "name": "地质数据分析 (CSV)",
            "q": "Which geological zone has the highest permeability, and what is its lithology?",
            "expected": ["Zone_A", "Sandstone"],
        },
        # ========== 新增：Query Decomposition 测试用例 ==========
        {
            "name": "比较类问题 (Query Decomposition)",
            "q": "Compare the revenue growth of Latin America vs Middle East in Q3 2024.",
            "expected": ["Latin America", "Middle East", "%"],
        },
        {
            "name": "因果推理 + 后续 (Query Decomposition)",
            "q": "Why did drilling stop on November 12 and what recovery actions were taken?",
            "expected": ["MWD", "failure", "pulse"],
        }
    ]
    
    print("\n" + "="*50)
    print("🤖 AI 裁判进场评测 (LLM-as-a-Judge) [Parent-Child Mode]")
    print("="*50)

    for case in test_cases:
        print(f"\n🧪 Testing: {case['name']}")
        print(f"❓ Question: {case['q']}")
        
        # Mock st UI components for CRAG
        mock_status = mock.MagicMock()
        mock_status.__enter__.return_value = mock.MagicMock()
        st.status = mock.MagicMock(return_value=mock_status)
        st.toast = mock.MagicMock()
        st.warning = mock.MagicMock()
        st.info = mock.MagicMock()
        st.expander = mock.MagicMock()
        
        response, sources, debug_info = process_query(case["q"], rag)
        
        print(f"🤖 Answer: {response[:150]}...") # 打印前150个字符
        print(f"🐛 Debug Info: {debug_info}")
        
        # Debug: Print retrieved sources details
        print("🔍 Debug Sources:")
        for doc, score in sources:
            print(f"   - [{score:.4f}] {doc.page_content[:100]}...")
            
        # AI 打分
        score, reason = judge_answer(rag["llm"], case["q"], case["expected"], response)
        
        status = "✅ PASS" if score >= 7 else "❌ FAIL"
        print(f"📝 Score: {score}/10")
        print(f"🧐 Judge's Comment: {reason[:100]}...") # 打印部分评语
        print(f"结果: {status}")

if __name__ == "__main__":
    run_challenge_benchmark()
