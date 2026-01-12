import os
import sys

# 强制绕过代理，解决 502 错误
os.environ["NO_PROXY"] = "localhost,127.0.0.1"

import pandas as pd
from langchain_community.document_loaders import (
    TextLoader, 
    DirectoryLoader, 
    PyPDFLoader, 
    Docx2txtLoader
)
from langchain_text_splitters import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from sentence_transformers import CrossEncoder
from langchain_core.documents import Document
from langchain_ollama import ChatOllama

# 直接使用 rank_bm25，不依赖 LangChain 的 BM25Retriever
from rank_bm25 import BM25Okapi

# 检查环境变量
if "OPENAI_API_KEY" not in os.environ:
    print("⚠️  警告: 未检测到 OPENAI_API_KEY 环境变量 (LLM 部分可能无法运行)。")

def load_excel_as_text(directory):
    """
    专门优化 Excel 加载：Table-to-Text
    """
    documents = []
    import glob
    excel_files = glob.glob(os.path.join(directory, "**/*.xlsx"), recursive=True)
    
    for file_path in excel_files:
        try:
            df = pd.read_excel(file_path)
            df = df.fillna("Unknown")
            for index, row in df.iterrows():
                content = (
                    f"Market Data Record: Company {row.get('Company', '')} "
                    f"is located in {row.get('Country', '')}. "
                    f"Its 2023 Revenue was {row.get('Revenue_2023_Billion_USD', '')} Billion USD. "
                    f"It has {row.get('Employees', '')} employees. "
                    f"Key service focus is {row.get('Key_Service', '')}."
                )
                doc = Document(
                    page_content=content, 
                    metadata={"source": file_path, "type": "excel_record"}
                )
                documents.append(doc)
            print(f"✅ 成功加载 Excel (转自然语言): {file_path} (行数: {len(df)})")
        except Exception as e:
            print(f"⚠️  Excel 加载失败 {file_path}: {e}")
            
    return documents

def load_word_with_structure(directory):
    """
    专门优化 Word 加载：保留文档结构 (Title + Content)
    对于合同，"Section 2. Compensation" 这样的标题对检索至关重要。
    """
    documents = []
    import glob
    from docx import Document as DocxDocument # 需要 pip install python-docx
    
    word_files = glob.glob(os.path.join(directory, "**/*.docx"), recursive=True)
    
    for file_path in word_files:
        try:
            doc_obj = DocxDocument(file_path)
            current_heading = "General"
            
            for para in doc_obj.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # 简单判断：如果是粗体或者字数较少，可能是标题
                # 这里假设 style.name 包含 'Heading' 是标题
                if 'Heading' in para.style.name:
                    current_heading = text
                else:
                    # 将标题拼接到内容前，增强语义
                    # 例如: "Section 2. Compensation: Client agrees to pay..."
                    enhanced_content = f"Document Section [{current_heading}]: {text}"
                    
                    doc = Document(
                        page_content=enhanced_content, 
                        metadata={"source": file_path, "type": "contract_clause", "section": current_heading}
                    )
                    documents.append(doc)
                    
            print(f"✅ 成功加载 Word (结构化): {file_path}")
        except Exception as e:
            print(f"⚠️  Word 加载失败 {file_path}: {e}")
            
    return documents

def run_rag_demo():
    print("--- 1. Data Processing (数据处理优化版) ---")
    
    docs = []
    
    # 1.1 加载 Markdown (使用 MarkdownHeaderTextSplitter 优化)
    # 不再用 TextLoader 傻读，而是读取内容后按标题切分
    import glob
    md_files = glob.glob(os.path.join("knowledge_base", "**/*.md"), recursive=True)
    for f in md_files:
        with open(f, 'r', encoding='utf-8') as file:
            md_content = file.read()
            
        # 定义切分规则：按 H1, H2, H3 切分，保留层级结构
        headers_to_split_on = [
            ("#", "Header 1"),
            ("##", "Header 2"),
            ("###", "Header 3"),
        ]
        markdown_splitter = MarkdownHeaderTextSplitter(headers_to_split_on=headers_to_split_on)
        md_header_splits = markdown_splitter.split_text(md_content)
        
        # 给每个切分出来的块加上 source 元数据
        for doc in md_header_splits:
            doc.metadata["source"] = f
            doc.metadata["type"] = "markdown_section"
            # 这一步很关键：把标题拼回正文，增强语义
            # MarkdownHeaderTextSplitter 会把标题放在 metadata 里，我们需要把它拿出来
            header_path = " > ".join([v for k, v in doc.metadata.items() if k.startswith("Header")])
            if header_path:
                doc.page_content = f"Section [{header_path}]:\n{doc.page_content}"
                
        docs.extend(md_header_splits)
        print(f"✅ 成功加载 Markdown (按标题切分): {f} (块数: {len(md_header_splits)})")
    
    # 1.2 加载 Word (改为结构化加载)
    # 移除旧的 Docx2txtLoader
    # loader_word = DirectoryLoader("downloads", glob="**/*.docx", loader_cls=Docx2txtLoader)
    # docs.extend(loader_word.load())
    docs.extend(load_word_with_structure("downloads"))
    
    # 1.3 加载 Excel (Table-to-Text)
    docs.extend(load_excel_as_text("downloads"))
    
    # 1.4 加载 PDF
    print("正在加载 PDF...")
    pdf_loader = DirectoryLoader("downloads", glob="**/*.pdf", loader_cls=PyPDFLoader)
    docs.extend(pdf_loader.load())
    
    print(f"总计加载文档片段: {len(docs)}")

    # 1.5 切分文档 (仅切分 PDF 和其他长文本，Markdown/Excel/Word 已经切好了)
    # 这里的逻辑需要调整：我们只对那些太长的 doc 进行二次切分
    final_docs = []
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    
    for doc in docs:
        # 如果是 PDF (type="pdf" 或未标记)，切分
        # 如果是 Markdown/Excel/Word (已经结构化处理过)，如果太长也切一下，但尽量保留完整
        if len(doc.page_content) > 1000:
             final_docs.extend(text_splitter.split_documents([doc]))
        else:
             final_docs.append(doc)
             
    splits = final_docs
    print(f"文档预处理完成，共 {len(splits)} 个 Chunks。")

    print("\n--- 2. Building Retrievers (手动构建混合检索) ---")
    
    # 2.1 向量检索 (Chroma)
    print("初始化 Vector Store...")
    # 关键修改：切换到支持多语言的 Embedding 模型
    # paraphrase-multilingual-MiniLM-L12-v2: 轻量级但支持50+种语言，中英文对齐效果远好于 all-MiniLM
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")
    vectorstore = Chroma.from_documents(documents=splits, embedding=embeddings)
    
    # 2.2 关键词检索 (BM25)
    print("初始化 BM25 Index...")
    # 对所有 chunk 进行分词
    tokenized_corpus = [doc.page_content.split(" ") for doc in splits]
    bm25 = BM25Okapi(tokenized_corpus)

    # 2.3 Reranker
    print("加载 Reranker 模型...")
    reranker = CrossEncoder("BAAI/bge-reranker-base")
    
    # 2.4 初始化本地 LLM (Ollama)
    print("初始化本地 LLM (Ollama)...")
    try:
        # 使用 qwen2.5:3b，显式指定 URL
        llm = ChatOllama(
            model="qwen2.5:3b",
            base_url="http://127.0.0.1:11434"
        )
        # 测试一下 LLM 是否可用
        print("   -> 测试 LLM 连接...")
        llm.invoke("Hi")
        print("   -> LLM 连接成功！")
        use_llm = True
    except Exception as e:
        print(f"⚠️  本地 LLM 初始化失败: {e}")
        print("   -> 将仅返回检索到的文档片段，不进行生成。")
        use_llm = False

    print("\n--- 3. Execution (执行检索 + 生成) ---")
    
    questions = [
        "Schlumberger 的 2023 年营收是多少？(Excel数据)",
        "合同中规定的钻井日费率(day rate)是多少？(Word数据)",
        "什么是水力压裂(Fracking)？(Markdown知识库)",
    ]

    for q in questions:
        print(f"\n" + "="*50)
        print(f"用户提问: {q}")
        print("-" * 50)
        
        # --- Step 1: 混合召回 (Hybrid Retrieval) - 采用分组召回策略 (Grouped Retrieval) ---
        print("🔍 1. 分组混合召回 (Grouped Retrieval)...")
        
        candidate_docs = []
        seen_content = set()
        
        # 定义分组策略：强制从每种文档类型里都捞一点出来
        # 这样就能避免大文件 (PDF) 淹没小文件 (Excel/Word)
        filters = [
            {"name": "Excel", "filter": {"type": "excel_record"}, "k": 5},
            {"name": "Word", "filter": {"type": "contract_clause"}, "k": 5},
            {"name": "General", "filter": None, "k": 10} # General 负责捞 PDF 和其他未分类的
        ]
        
        for f in filters:
            kwargs = {"k": f["k"]}
            if f["filter"]:
                kwargs["filter"] = f["filter"]
                
            try:
                sub_docs = vectorstore.similarity_search(q, **kwargs)
                print(f"   - [{f['name']}] 向量召回: {len(sub_docs)} 个")
                
                for doc in sub_docs:
                    if doc.page_content not in seen_content:
                        # 简单的意图判断加成
                        if f["name"].lower() in q.lower() or (f["name"]=="Word" and "合同" in q):
                             doc.metadata["boost"] = True
                        candidate_docs.append(doc)
                        seen_content.add(doc.page_content)
            except Exception as e:
                print(f"   - [{f['name']}] 召回失败: {e}")

        # BM25 补充 (关键词匹配，防止向量模型“偏科”)
        tokenized_query = q.split(" ")
        bm25_top_n = bm25.get_top_n(tokenized_query, splits, n=10)
        for doc in bm25_top_n:
            if doc.page_content not in seen_content:
                candidate_docs.append(doc)
                seen_content.add(doc.page_content)
                
        print(f"   -> 候选文档总数: {len(candidate_docs)}")

        # --- Step 2: 重排序 (Rerank) ---
        print("🔍 2. 重排序 (Reranking)...")
        pairs = [[q, doc.page_content] for doc in candidate_docs]
        scores = reranker.predict(pairs)
        
        doc_score_pairs = list(zip(candidate_docs, scores))
        
        final_pairs = []
        for doc, score in doc_score_pairs:
            final_score = score
            if doc.metadata.get("boost"):
                print(f"   [Boost] 发现用户意图匹配文档: {os.path.basename(doc.metadata.get('source', ''))}")
                final_score += 0.5 
            final_pairs.append((doc, final_score))
            
        final_pairs.sort(key=lambda x: x[1], reverse=True)
        
        # 选出 Top 3
        top_docs = [doc for doc, score in final_pairs[:3]]
        
        # 展示检索结果
        for i, (doc, score) in enumerate(final_pairs[:3]):
            print(f"\n[排名 #{i+1} | 得分: {score:.4f}]:")
            short_source = os.path.basename(doc.metadata.get('source', 'unknown'))
            print(f"来源: {short_source}")
            print(f"内容: {doc.page_content[:200].replace(chr(10), ' ')}...")
            
        # --- Step 3: LLM 生成 (Generation) ---
        if use_llm:
            print("\n🤖 3. LLM 生成回答...")
            context = "\n\n".join([d.page_content for d in top_docs])
            prompt = f"""基于以下参考资料回答问题。如果资料中没有答案，请直接说不知道。

参考资料:
{context}

问题: {q}

回答:"""
            response = llm.invoke(prompt)
            print(f"\n✅ AI 回答:\n{response.content}")
        else:
            print("\n(跳过生成步骤，仅展示检索结果)")

    print("\n✅ 演示结束。")

if __name__ == "__main__":
    run_rag_demo()
