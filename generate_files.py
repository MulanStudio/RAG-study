import pandas as pd
from docx import Document
import os

# 1. Generate Excel File
excel_data = {
    "Company": ["Schlumberger (SLB)", "Halliburton", "Baker Hughes", "COSL", "Sinopec SSC"],
    "Country": ["USA/France", "USA", "USA", "China", "China"],
    "Revenue_2023_Billion_USD": [33.1, 23.0, 25.5, 6.2, 9.8],
    "Employees": [111000, 48000, 57000, 16000, 68000],
    "Key_Service": ["Wireline Logging", "Fracking", "Turbines", "Offshore Drilling", "Onshore Engineering"]
}

df = pd.DataFrame(excel_data)
excel_path = "downloads/Global_Oilfield_Market_Data.xlsx"
df.to_excel(excel_path, index=False)
print(f"Generated Excel: {excel_path}")

# 2. Generate Word File
doc = Document()
doc.add_heading('Master Service Agreement for Drilling Operations', 0)

doc.add_paragraph('This Master Service Agreement ("Agreement") is entered into by and between:')
doc.add_paragraph('Client: Global Energy Corp', style='List Bullet')
doc.add_paragraph('Contractor: Advanced Drilling Services Ltd.', style='List Bullet')

doc.add_heading('1. Scope of Work', level=1)
doc.add_paragraph(
    "The Contractor shall provide drilling rigs, personnel, and equipment necessary "
    "to drill three (3) exploratory wells in the Permian Basin region."
)

doc.add_heading('2. Compensation', level=1)
doc.add_paragraph(
    "Client agrees to pay Contractor a day rate of $25,000 USD per rig per day. "
    "Mobilization fees shall be paid separately upon arrival at the well site."
)

doc.add_heading('3. Safety and Compliance', level=1)
doc.add_paragraph(
    "Both parties agree to adhere to all API (American Petroleum Institute) standards "
    "and local environmental regulations."
)

doc_path = "downloads/Service_Contract_Template.docx"
doc.save(doc_path)
print(f"Generated Word: {doc_path}")

